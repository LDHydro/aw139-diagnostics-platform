"""
Document parsers.

Governing documents arrive as PDF, Word, HTML or Markdown.  Every parser
emits the same ``Block`` stream: a piece of text, the page it came from, and
whether it looks like a heading.  Keeping page numbers is what lets an answer
say "p. 51" instead of "somewhere in the manual".
"""

from __future__ import annotations

import html
import logging
import re
from dataclasses import dataclass, field
from pathlib import Path

log = logging.getLogger(__name__)

SUPPORTED_SUFFIXES = {
    ".pdf", ".docx", ".md", ".markdown", ".txt", ".html", ".htm", ".xml",
}


class ParseError(RuntimeError):
    """The document could not be read."""


@dataclass
class Block:
    """One paragraph, list item or heading."""

    text: str
    page: int | None = None
    heading_level: int | None = None  # 1-6 when this block is itself a heading

    @property
    def is_heading(self) -> bool:
        return self.heading_level is not None


@dataclass
class ParsedDocument:
    blocks: list[Block] = field(default_factory=list)
    page_count: int = 0
    meta: dict = field(default_factory=dict)

    @property
    def text(self) -> str:
        return "\n\n".join(b.text for b in self.blocks)


# ----------------------------------------------------------------------
# PDF
# ----------------------------------------------------------------------

def parse_pdf(path: Path) -> ParsedDocument:
    """
    Extract text from a PDF, inferring headings from typography.

    PDFs carry no semantic structure, so we use the two signals that
    actually survive export: a line set noticeably larger than body text,
    or a bold short line.  Combined with the numbering regex in the chunker
    this recovers the section hierarchy of most manuals.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ParseError("PyMuPDF is required to read PDF files") from exc

    doc = fitz.open(str(path))
    blocks: list[Block] = []

    # Body-text size, estimated across a sample of pages.
    sizes: list[float] = []
    sample = min(len(doc), 12)
    for page_index in range(sample):
        for block in doc[page_index].get_text("dict").get("blocks", []):
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    if span.get("text", "").strip():
                        sizes.append(round(span.get("size", 0.0), 1))
    body_size = _mode(sizes) if sizes else 10.0

    for page_index in range(len(doc)):
        page = doc[page_index]
        page_number = page_index + 1
        for block in page.get_text("dict").get("blocks", []):
            if block.get("type") != 0:  # skip images
                continue
            line_texts: list[str] = []
            max_size = 0.0
            bold = False
            for line in block.get("lines", []):
                span_texts = []
                for span in line.get("spans", []):
                    text = span.get("text", "")
                    if not text.strip():
                        continue
                    span_texts.append(text)
                    max_size = max(max_size, span.get("size", 0.0))
                    # Bit 4 of the span flags marks a bold face.
                    if span.get("flags", 0) & 2 ** 4:
                        bold = True
                    font = str(span.get("font", "")).lower()
                    if "bold" in font or "black" in font:
                        bold = True
                if span_texts:
                    line_texts.append("".join(span_texts).strip())

            text = " ".join(t for t in line_texts if t).strip()
            if not text:
                continue

            heading_level = None
            if max_size >= body_size * 1.35:
                heading_level = 1
            elif max_size >= body_size * 1.15:
                heading_level = 2
            elif bold and len(text) < 120 and len(line_texts) <= 2:
                heading_level = 3

            blocks.append(Block(text=text, page=page_number, heading_level=heading_level))

    result = ParsedDocument(
        blocks=blocks,
        page_count=len(doc),
        meta={k: v for k, v in (doc.metadata or {}).items() if v},
    )
    doc.close()
    return result


def _mode(values: list[float]) -> float:
    counts: dict[float, int] = {}
    for value in values:
        counts[value] = counts.get(value, 0) + 1
    return max(counts.items(), key=lambda kv: kv[1])[0]


# ----------------------------------------------------------------------
# Word
# ----------------------------------------------------------------------

def parse_docx(path: Path) -> ParsedDocument:
    """Word documents carry real heading styles - use them directly."""
    try:
        import docx  # python-docx
    except ImportError as exc:  # pragma: no cover - optional dependency
        raise ParseError("python-docx is required to read .docx files") from exc

    document = docx.Document(str(path))
    blocks: list[Block] = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        style = (paragraph.style.name if paragraph.style else "") or ""
        heading_level = None
        match = re.match(r"Heading (\d)", style, re.IGNORECASE)
        if match:
            heading_level = int(match.group(1))
        elif style.lower() in {"title", "subtitle"}:
            heading_level = 1
        blocks.append(Block(text=text, heading_level=heading_level))

    # Tables often hold the interval/limit data in a maintenance schedule,
    # so flatten them rather than dropping them.
    for table in document.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            if any(cells):
                rows.append(" | ".join(cells))
        if rows:
            blocks.append(Block(text="\n".join(rows)))

    core = document.core_properties
    meta = {
        "title": core.title or "",
        "author": core.author or "",
        "revision": str(core.revision or ""),
    }
    return ParsedDocument(blocks=blocks, page_count=0, meta={k: v for k, v in meta.items() if v})


# ----------------------------------------------------------------------
# Markdown / plain text / HTML / XML
# ----------------------------------------------------------------------

_MD_HEADING = re.compile(r"^(#{1,6})\s+(.*)$")


def parse_markdown(path: Path) -> ParsedDocument:
    blocks: list[Block] = []
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            text = "\n".join(buffer).strip()
            if text:
                blocks.append(Block(text=text))
            buffer.clear()

    for line in path.read_text(encoding="utf-8", errors="replace").splitlines():
        match = _MD_HEADING.match(line)
        if match:
            flush()
            blocks.append(Block(text=match.group(2).strip(), heading_level=len(match.group(1))))
        elif not line.strip():
            flush()
        else:
            buffer.append(line)
    flush()
    return ParsedDocument(blocks=blocks)


def parse_text(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    blocks = [
        Block(text=part.strip())
        for part in re.split(r"\n\s*\n", raw)
        if part.strip()
    ]
    return ParsedDocument(blocks=blocks)


_TAG_RE = re.compile(r"<[^>]+>")
_SCRIPT_RE = re.compile(r"<(script|style)[^>]*>.*?</\1>", re.IGNORECASE | re.DOTALL)
_HTML_HEADING_RE = re.compile(r"<h([1-6])[^>]*>(.*?)</h\1>", re.IGNORECASE | re.DOTALL)
_BLOCK_SPLIT_RE = re.compile(
    r"</(?:p|div|li|tr|section|article|h[1-6])>", re.IGNORECASE
)


def parse_html(path: Path) -> ParsedDocument:
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = _SCRIPT_RE.sub(" ", raw)

    blocks: list[Block] = []
    for fragment in _BLOCK_SPLIT_RE.split(raw):
        heading = _HTML_HEADING_RE.search(fragment + "</h1>")
        level = None
        if heading and heading.start() < 200:
            level = int(heading.group(1))
        text = html.unescape(_TAG_RE.sub(" ", fragment))
        text = re.sub(r"\s+", " ", text).strip()
        if text:
            blocks.append(Block(text=text, heading_level=level))
    return ParsedDocument(blocks=blocks)


def parse_xml(path: Path) -> ParsedDocument:
    """
    Flatten XML to text.

    Used for S1000D data modules and similar structured manuals, where the
    element names carry little meaning once the text is extracted.
    """
    raw = path.read_text(encoding="utf-8", errors="replace")
    raw = re.sub(r"<\?xml[^>]*\?>", " ", raw)
    text = html.unescape(_TAG_RE.sub("\n", raw))
    blocks = [
        Block(text=re.sub(r"\s+", " ", part).strip())
        for part in re.split(r"\n\s*\n", text)
        if part.strip()
    ]
    return ParsedDocument(blocks=[b for b in blocks if len(b.text) > 2])


_PARSERS = {
    ".pdf": parse_pdf,
    ".docx": parse_docx,
    ".md": parse_markdown,
    ".markdown": parse_markdown,
    ".txt": parse_text,
    ".html": parse_html,
    ".htm": parse_html,
    ".xml": parse_xml,
}


def parse(path: Path) -> ParsedDocument:
    """Dispatch on file extension."""
    suffix = path.suffix.lower()
    parser = _PARSERS.get(suffix)
    if parser is None:
        raise ParseError(
            f"unsupported file type '{suffix}'. Supported: "
            + ", ".join(sorted(SUPPORTED_SUFFIXES))
        )
    if not path.is_file():
        raise ParseError(f"not a file: {path}")
    return parser(path)
